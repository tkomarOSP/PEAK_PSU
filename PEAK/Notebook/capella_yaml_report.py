
"""
capella_yaml_report.py (v3)
---------------------------
Extensible report builder for Capella-style YAML exports.

Reports:
- Capability/CapabilityRealization × Functional Chain matrix
- Unlinked Capability-like items
- Requirement → Objects relation table (handles "relations" → link objects with source/target)
- Capability-centric requirement rollup (Direct → FC → Function → Component)
"""

from __future__ import annotations
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple
import os, re, yaml, pandas as pd

__all__ = ["CapellaYamlReport"]


class CapellaYamlReport:
    # ---------------- defaults ----------------
    DEFAULT_ID_KEYS = ("primary_uuid", "uuid", "id", "ref_uuid", "xid")
    DEFAULT_NAME_KEYS = ("long_name", "name", "title", "label")

    DEFAULT_CAPABILITY_TOKENS = frozenset({"capability", "operationalcapability", "businesscapability"})
    DEFAULT_CAPABILITY_REALIZATION_TOKENS = frozenset({"capabilityrealization"})
    DEFAULT_FUNCTIONAL_CHAIN_TOKENS = frozenset({"functionalchain"})
    DEFAULT_FUNCTION_TOKENS = frozenset({"function"})
    DEFAULT_REQUIREMENT_TOKENS = frozenset({"requirement"})
    DEFAULT_COMPONENT_TOKENS = frozenset({"component","logicalcomponent","physicalcomponent","systemcomponent"})

    DEFAULT_CAPABILITY_CONTAINER_TOKENS = frozenset({"capabilities", "capabilitypkg", "capabilitypackage"})
    DEFAULT_CAPABILITY_REALIZATION_CONTAINER_TOKENS = frozenset({"capabilityrealizations","capabilityrealizationpkg","capabilityrealizationpackage"})
    DEFAULT_REQUIREMENT_CONTAINER_TOKENS = frozenset({"requirements","requirementpkg","requirementpackage"})
    DEFAULT_COMPONENT_CONTAINER_TOKENS = frozenset({"components","logicalcomponents","physicalcomponents","componentpkg","componentpackage"})

    DEFAULT_CAP_TO_FC_KEYS = (
        "involved_functional_chains","functional_chains","involvedFunctionalChains",
        "related_functional_chains","functionalChains","allocated_functional_chains"
    )
    DEFAULT_FC_TO_CAP_KEYS = (
        "capabilities","involved_capabilities","involvedCapabilities","related_capabilities"
    )

    DEFAULT_REQ_REL_KEYS = (
        "target","targets","links","link",
        "satisfied_elements","satisfies",
        "verified_elements","verifies",
        "refined_elements","refines",
        "traced_to","trace","traceability",
        "allocated_to","allocated_from",
        "related_elements","involve"
    )

    def __init__(
        self, yaml_paths: List[str], *,
        capability_tokens: Optional[Set[str]] = None,
        capability_realization_tokens: Optional[Set[str]] = None,
        functional_chain_tokens: Optional[Set[str]] = None,
        function_tokens: Optional[Set[str]] = None,
        requirement_tokens: Optional[Set[str]] = None,
        component_tokens: Optional[Set[str]] = None,
        capability_container_tokens: Optional[Set[str]] = None,
        capability_realization_container_tokens: Optional[Set[str]] = None,
        requirement_container_tokens: Optional[Set[str]] = None,
        component_container_tokens: Optional[Set[str]] = None,
        id_keys: Optional[Iterable[str]] = None,
        name_keys: Optional[Iterable[str]] = None,
        cap_to_fc_keys: Optional[Iterable[str]] = None,
        fc_to_cap_keys: Optional[Iterable[str]] = None,
    ) -> None:
        self.yaml_paths = list(yaml_paths)

        self.capability_tokens = set(capability_tokens or self.DEFAULT_CAPABILITY_TOKENS)
        self.capability_realization_tokens = set(capability_realization_tokens or self.DEFAULT_CAPABILITY_REALIZATION_TOKENS)
        self.functional_chain_tokens = set(functional_chain_tokens or self.DEFAULT_FUNCTIONAL_CHAIN_TOKENS)
        self.function_tokens = set(function_tokens or self.DEFAULT_FUNCTION_TOKENS)
        self.requirement_tokens = set(requirement_tokens or self.DEFAULT_REQUIREMENT_TOKENS)
        self.component_tokens = set(component_tokens or self.DEFAULT_COMPONENT_TOKENS)

        self.capability_container_tokens = set(capability_container_tokens or self.DEFAULT_CAPABILITY_CONTAINER_TOKENS)
        self.capability_realization_container_tokens = set(capability_realization_container_tokens or self.DEFAULT_CAPABILITY_REALIZATION_CONTAINER_TOKENS)
        self.requirement_container_tokens = set(requirement_container_tokens or self.DEFAULT_REQUIREMENT_CONTAINER_TOKENS)
        self.component_container_tokens = set(component_container_tokens or self.DEFAULT_COMPONENT_CONTAINER_TOKENS)

        self.id_keys = tuple(id_keys or self.DEFAULT_ID_KEYS)
        self.name_keys = tuple(name_keys or self.DEFAULT_NAME_KEYS)

        self.cap_to_fc_keys = tuple(cap_to_fc_keys or self.DEFAULT_CAP_TO_FC_KEYS)
        self.fc_to_cap_keys = tuple(fc_to_cap_keys or self.DEFAULT_FC_TO_CAP_KEYS)

        # caches
        self._raw: Any = None
        self._objects: List[Dict[str, Any]] = []
        self._caps: Dict[str, Dict[str, Any]] = {}
        self._crs: Dict[str, Dict[str, Any]] = {}
        self._caplikes: Dict[str, Dict[str, Any]] = {}
        self._fcs: Dict[str, Dict[str, Any]] = {}
        self._funcs: Dict[str, Dict[str, Any]] = {}
        self._requirements: Dict[str, Dict[str, Any]] = {}
        self._components: Dict[str, Dict[str, Any]] = {}
        self._relations_by_uuid: Dict[str, Dict[str, Any]] = {}

        self._comp_to_fn: Dict[str, Set[str]] = {}
        self._fn_to_comp: Dict[str, Set[str]] = {}

        self.reload()

    # -------------- public API --------------
    def reload(self) -> None:
        docs = []
        for p in self.yaml_paths:
            with open(p, "r", encoding="utf-8") as f:
                docs.append(yaml.safe_load(f))
        self._raw = docs[0] if len(docs) == 1 else {"docs": docs}
        self._objects = self._index_objects(self._raw)

        self._caps = self._collect_by_type(self._objects, self.capability_tokens)
        self._crs  = self._collect_by_type(self._objects, self.capability_realization_tokens)
        self._caplikes = {**{k:{**v,"kind":"Capability"} for k,v in self._caps.items()},
                          **{k:{**v,"kind":"CapabilityRealization"} for k,v in self._crs.items()}}
        self._fcs  = self._collect_by_type(self._objects, self.functional_chain_tokens)
        self._funcs= self._collect_by_type(self._objects, self.function_tokens)

        self._apply_container_caps()
        self._requirements = self._collect_by_type(self._objects, self.requirement_tokens)
        self._components   = self._collect_by_type(self._objects, self.component_tokens)
        self._apply_container_reqs_components()

        # v3: relation objects index + heuristic requirements (text + relations)
        self._relations_by_uuid = {}
        for o in self._objects:
            if isinstance(o, dict) and "primary_uuid" in o and "source" in o and "target" in o:
                self._relations_by_uuid[o["primary_uuid"]] = o
        for o in self._objects:
            if isinstance(o, dict) and o.get("primary_uuid") and isinstance(o.get("relations"), list) and isinstance(o.get("text"), str):
                oid = o["primary_uuid"]
                if oid not in self._requirements:
                    nm = o.get("name") or o.get("long_name") or o.get("title") or oid
                    self._requirements[oid] = {"id": oid, "name": nm, "node": o}

        self._build_function_component_adjacency()

    def capability_fc_matrix(self) -> pd.DataFrame:
        cap_to_fc, _ = self._caplike_to_fc_links()
        cap_rows = sorted(((cid, self._caplikes[cid]["name"], self._caplikes[cid]["kind"]) for cid in self._caplikes), key=lambda t:(t[2], t[1].lower()))
        fc_cols  = sorted(((fid, self._fcs[fid]["name"]) for fid in self._fcs), key=lambda t:t[1].lower())
        rows = []
        for cid, cname, ckind in cap_rows:
            row = {"Capability": cname, "Kind": ckind}
            linked = cap_to_fc.get(cid, set())
            for fid, fname in fc_cols:
                row[fname] = "✓" if fid in linked else ""
            rows.append(row)
        return pd.DataFrame.from_records(rows, columns=["Capability","Kind"]+[fname for _,fname in fc_cols])

    def capabilities_without_fc(self) -> pd.DataFrame:
        cap_to_fc, _ = self._caplike_to_fc_links()
        rows = []
        for cid, cent in sorted(self._caplikes.items(), key=lambda kv: kv[1]["name"].lower()):
            if not cap_to_fc.get(cid):
                rows.append({"Capability": cent["name"], "Kind": cent["kind"], "ID": cid})
        return pd.DataFrame.from_records(rows, columns=["Capability","Kind","ID"])

    def requirements_relation_table(self) -> pd.DataFrame:
        req_to_objs = self._requirement_relations()
        rows = []
        for rid, r in self._requirements.items():
            rname = r["name"]
            stmt, rat = self._requirement_text_fields(r["node"])
            targets = sorted(req_to_objs.get(rid, set()))
            if not targets:
                rows.append({"Requirement": rname, "RequirementID": rid, "Statement": stmt, "Rationale": rat,
                             "RelatedObject": "", "RelatedID": "", "RelatedKind": ""})
            else:
                for oid in targets:
                    rows.append({"Requirement": rname, "RequirementID": rid, "Statement": stmt, "Rationale": rat,
                                 "RelatedObject": self._name_for_id(oid), "RelatedID": oid, "RelatedKind": self._object_kind(oid)})
        return pd.DataFrame.from_records(rows, columns=["Requirement","RequirementID","Statement","Rationale","RelatedObject","RelatedID","RelatedKind"])

    def capability_requirements_report(self) -> pd.DataFrame:
        req_to_objs = self._requirement_relations()
        # reverse index
        obj_to_reqs: Dict[str, Set[str]] = {}
        for rqid, objset in req_to_objs.items():
            for oid in objset:
                obj_to_reqs.setdefault(oid, set()).add(rqid)

        cap_to_fc, _ = self._caplike_to_fc_links()
        fc_to_fn = {fid: self._fc_functions(self._fcs[fid]["node"], set(self._funcs.keys())) for fid in self._fcs}
        fn_to_comp = self._fn_to_comp

        rows = []
        for cid, cent in sorted(self._caplikes.items(), key=lambda kv: kv[1]["name"].lower()):
            cname, ckind = cent["name"], cent["kind"]
            direct_objs = {cid}
            fc_objs = set(cap_to_fc.get(cid, set()))
            fn_objs = set().union(*(fc_to_fn.get(fid, set()) for fid in fc_objs)) if fc_objs else set()
            comp_objs: Set[str] = set()
            for fnid in fn_objs:
                comp_objs |= fn_to_comp.get(fnid, set())

            def emit(bucket, via):
                for oid in sorted(bucket):
                    for rqid in sorted(obj_to_reqs.get(oid, set())):
                        rnode = self._requirements[rqid]
                        stmt, rat = self._requirement_text_fields(rnode["node"])
                        rows.append({
                            "Capability": cname, "Kind": ckind,
                            "Requirement": rnode["name"], "RequirementID": rqid,
                            "Statement": stmt, "Rationale": rat,
                            "MatchedVia": via, "MatchedObject": self._name_for_id(oid), "MatchedID": oid
                        })

            emit(direct_objs, "Direct")
            emit(fc_objs, "FunctionalChain")
            emit(fn_objs, "Function")
            emit(comp_objs, "Component")

        cols = ["Capability","Kind","Requirement","RequirementID","Statement","Rationale","MatchedVia","MatchedObject","MatchedID"]
        df = pd.DataFrame.from_records(rows, columns=cols)
        if not df.empty:
            df = df.sort_values(["Capability","Kind","Requirement","MatchedVia","MatchedObject"]).reset_index(drop=True)
        return df

    def save_csv(self, df: pd.DataFrame, path: str) -> str:
        abspath = os.path.abspath(path); df.to_csv(abspath, index=False); return abspath

    # -------------- internals --------------
    @staticmethod
    def _type_contains(t: Optional[str], tokens: Set[str]) -> bool:
        if not t or not isinstance(t, str): return False
        tt = t.replace("::"," ").replace("."," ").replace("_","").replace(" ","").lower()
        return any(tok in tt for tok in tokens)

    @staticmethod
    def _normalize_key(s: str) -> str:
        return s.replace("_","").replace(" ","").lower()

    def _index_objects(self, y: Any) -> List[Dict[str, Any]]:
        if isinstance(y, dict):
            if "model" in y and isinstance(y["model"], dict) and isinstance(y["model"].get("objects"), list):
                return y["model"]["objects"]
            return [n for n, _ in self._walk(y) if isinstance(n, dict) and "type" in n]
        elif isinstance(y, list):
            return [o for o in y if isinstance(o, dict)]
        return []

    @staticmethod
    def _walk(node: Any, parent_key: Optional[str] = None):
        if isinstance(node, dict):
            yield node, parent_key
            for k, v in node.items():
                for child, p in CapellaYamlReport._walk(v, k):
                    yield child, p
        elif isinstance(node, list):
            for item in node:
                for child, p in CapellaYamlReport._walk(item, parent_key):
                    yield child, p

    def _get_first(self, d: Dict[str, Any], keys: Iterable[str]) -> Optional[Any]:
        for k in keys:
            if isinstance(d, dict) and k in d:
                return d[k]
        return None

    def _extract_id(self, obj: Any) -> Optional[str]:
        if obj is None: return None
        if isinstance(obj, str): return obj.strip()
        if isinstance(obj, dict):
            v = self._get_first(obj, self.id_keys)
            if isinstance(v, str): return v.strip()
            for k in ("ref","ref_uuid","uuid","id"):
                vv = obj.get(k)
                if isinstance(vv, str): return vv.strip()
                if isinstance(vv, dict):
                    vv2 = self._get_first(vv, ("uuid","id","ref_uuid"))
                    if isinstance(vv2, str): return vv2.strip()
        return None

    def _extract_name(self, obj: Any) -> Optional[str]:
        if obj is None: return None
        if isinstance(obj, str): return obj.strip()
        if isinstance(obj, dict):
            v = self._get_first(obj, self.name_keys)
            if isinstance(v, str): return v.strip()
        return None

    def _collect_by_type(self, objs: List[Dict[str, Any]], tokens: Set[str]) -> Dict[str, Dict[str, Any]]:
        out: Dict[str, Dict[str, Any]] = {}
        for o in objs:
            if self._type_contains(o.get("type"), tokens):
                oid = self._extract_id(o) or f"NOID::{id(o)}"
                nm  = self._extract_name(o) or oid
                out[oid] = {"id": oid, "name": nm, "node": o}
        return out

    def _apply_container_caps(self) -> None:
        # CR containers
        for o, parent_key in self._walk(self._objects):
            if not isinstance(o, dict) or not parent_key: continue
            pk = self._normalize_key(parent_key)
            if any(tok in pk for tok in self.capability_realization_container_tokens):
                oid = self._extract_id(o) or f"NOID::{id(o)}"
                if oid not in self._crs:
                    nm = self._extract_name(o) or oid
                    self._crs[oid] = {"id": oid, "name": nm, "node": o}
                    self._caplikes[oid] = {"id": oid, "name": nm, "kind": "CapabilityRealization", "node": o}
        # Capability containers
        for o, parent_key in self._walk(self._objects):
            if not isinstance(o, dict) or not parent_key: continue
            pk = self._normalize_key(parent_key)
            if any(tok in pk for tok in self.capability_container_tokens):
                oid = self._extract_id(o) or f"NOID::{id(o)}"
                if oid not in self._caps and oid not in self._crs:
                    nm = self._extract_name(o) or oid
                    self._caps[oid] = {"id": oid, "name": nm, "node": o}
                    self._caplikes[oid] = {"id": oid, "name": nm, "kind": "Capability", "node": o}

    def _apply_container_reqs_components(self) -> None:
        for o, parent_key in self._walk(self._objects):
            if not isinstance(o, dict) or not parent_key: continue
            pk = self._normalize_key(parent_key)
            if any(tok in pk for tok in self.requirement_container_tokens):
                oid = self._extract_id(o) or f"NOID::{id(o)}"
                if oid not in self._requirements:
                    nm = self._extract_name(o) or oid
                    self._requirements[oid] = {"id": oid, "name": nm, "node": o}
            if any(tok in pk for tok in self.component_container_tokens):
                oid = self._extract_id(o) or f"NOID::{id(o)}"
                if oid not in self._components:
                    nm = self._extract_name(o) or oid
                    self._components[oid] = {"id": oid, "name": nm, "node": o}

    def _build_function_component_adjacency(self) -> None:
        self._comp_to_fn = {cid: set() for cid in self._components}
        self._fn_to_comp = {fid: set() for fid in self._funcs}
        # components referencing functions
        for cid, comp in self._components.items():
            node = comp["node"]
            for rid in self._iter_all_refs(node):
                if rid in self._funcs:
                    self._comp_to_fn[cid].add(rid)
                    self._fn_to_comp.setdefault(rid, set()).add(cid)
        # functions referencing components
        for fid, fn in self._funcs.items():
            node = fn["node"]
            for rid in self._iter_all_refs(node):
                if rid in self._components:
                    self._fn_to_comp.setdefault(fid, set()).add(rid)
                    self._comp_to_fn.setdefault(rid, set()).add(fid)
        # augment with space-key variants
        extra_keys = ("allocated from", "allocated to", "functions allocated to", "functions allocated from")
        for fid, fn in self._funcs.items():
            node = fn["node"]
            for k in extra_keys:
                entries = node.get(k); 
                if entries is None: continue
                if not isinstance(entries, list): entries = [entries]
                for entry in entries:
                    oid = self._extract_id(entry if isinstance(entry, dict) else {"ref_uuid": entry})
                    if oid and oid in self._components:
                        self._fn_to_comp.setdefault(fid, set()).add(oid)
                        self._comp_to_fn.setdefault(oid, set()).add(fid)
        for cid, comp in self._components.items():
            node = comp["node"]
            for k in extra_keys:
                entries = node.get(k); 
                if entries is None: continue
                if not isinstance(entries, list): entries = [entries]
                for entry in entries:
                    oid = self._extract_id(entry if isinstance(entry, dict) else {"ref_uuid": entry})
                    if oid and oid in self._funcs:
                        self._comp_to_fn.setdefault(cid, set()).add(oid)
                        self._fn_to_comp.setdefault(oid, set()).add(cid)

    # link discovery
    def _iter_all_refs(self, node: Any):
        if isinstance(node, dict):
            for k, v in node.items():
                if isinstance(v, (str, dict)):
                    rid = self._extract_id({k: v} if not isinstance(v, dict) else v)
                    if isinstance(rid, str):
                        yield rid
            for v in node.values():
                yield from self._iter_all_refs(v)
        elif isinstance(node, list):
            for item in node:
                yield from self._iter_all_refs(item)

    def _gather_specific_refs(self, obj: Dict[str, Any], keys: Iterable[str]) -> Set[str]:
        out: Set[str] = set()
        for k in keys:
            if isinstance(obj, dict) and k in obj:
                entries = obj[k] if isinstance(obj[k], list) else [obj[k]]
                for entry in entries:
                    rid = self._extract_id(entry if isinstance(entry, dict) else {"ref_uuid": entry})
                    if rid: out.add(rid)
        return out

    def _fc_functions(self, fc_node: Dict[str, Any], function_ids: Set[str]) -> Set[str]:
        fn_ids: Set[str] = set()
        involve = fc_node.get("involve")
        entries = involve if isinstance(involve, list) else ([involve] if involve is not None else [])
        for entry in entries:
            if isinstance(entry, dict):
                if self._type_contains(entry.get("type",""), self.function_tokens):
                    rid = self._extract_id(entry)
                    if rid and rid in function_ids:
                        fn_ids.add(rid)
        for rid in self._iter_all_refs(fc_node):
            if rid in function_ids:
                fn_ids.add(rid)
        return fn_ids

    def _caplike_functions(self, cap_node: Dict[str, Any], function_ids: Set[str]) -> Set[str]:
        fn_ids: Set[str] = set()
        for key in ("functions","owned_functions","allocated_functions","involve","functions owned"):
            entries = cap_node.get(key)
            entries = entries if isinstance(entries, list) else ([entries] if entries is not None else [])
            for entry in entries:
                rid = self._extract_id(entry if isinstance(entry, dict) else {"ref_uuid": entry})
                if rid and rid in function_ids:
                    fn_ids.add(rid)
        for rid in self._iter_all_refs(cap_node):
            if rid in function_ids:
                fn_ids.add(rid)
        return fn_ids

    def _caplike_to_fc_links(self) -> Tuple[Dict[str, Set[str]], Dict[str, Set[str]]]:
        cap_ids = set(self._caplikes.keys())
        fc_ids  = set(self._fcs.keys())
        fn_ids  = set(self._funcs.keys())

        cap_to_fc: Dict[str, Set[str]] = {cid: set() for cid in cap_ids}
        fc_to_cap: Dict[str, Set[str]] = {fid: set() for fid in fc_ids}

        for cid in cap_ids:
            node = self._caplikes[cid]["node"]
            direct = self._gather_specific_refs(node, self.cap_to_fc_keys)
            for rid in self._iter_all_refs(node):
                if rid in fc_ids: direct.add(rid)
            cap_to_fc[cid] |= {rid for rid in direct if rid in fc_ids}

        for fid in fc_ids:
            node = self._fcs[fid]["node"]
            direct = self._gather_specific_refs(node, self.fc_to_cap_keys)
            for rid in self._iter_all_refs(node):
                if rid in cap_ids: direct.add(rid)
            fc_to_cap[fid] |= {rid for rid in direct if rid in cap_ids}

        fc_fn_map = {fid: self._fc_functions(self._fcs[fid]["node"], fn_ids) for fid in fc_ids}
        cap_fn_map= {cid: self._caplike_functions(self._caplikes[cid]["node"], fn_ids) for cid in cap_ids}

        for cid in cap_ids:
            for fid in fc_ids:
                if cap_fn_map[cid] & fc_fn_map[fid]:
                    cap_to_fc[cid].add(fid); fc_to_cap[fid].add(cid)

        for cid, fset in list(cap_to_fc.items()):
            for fid in fset: fc_to_cap.setdefault(fid, set()).add(cid)
        for fid, cset in list(fc_to_cap.items()):
            for cid in cset: cap_to_fc.setdefault(cid, set()).add(fid)

        return cap_to_fc, fc_to_cap

    def _object_kind(self, oid: str) -> str:
        if oid in self._caps: return "Capability"
        if oid in self._crs: return "CapabilityRealization"
        if oid in self._fcs: return "FunctionalChain"
        if oid in self._funcs: return "Function"
        if oid in self._components: return "Component"
        return "Object"

    def _name_for_id(self, oid: str) -> str:
        if oid in self._caplikes: return self._caplikes[oid]["name"]
        if oid in self._fcs: return self._fcs[oid]["name"]
        if oid in self._funcs: return self._funcs[oid]["name"]
        if oid in self._components: return self._components[oid]["name"]
        req = self._requirements.get(oid)
        if req: return req["name"]
        return oid

    # requirement helpers
    _MKUP_RE = re.compile(r"Markup\((['\"])(.*)\1\)", re.S)
    def _unwrap_markup(self, s: Any) -> str:
        if not isinstance(s, str): return str(s) if s is not None else ""
        m = self._MKUP_RE.fullmatch(s)
        return m.group(2) if m else s

    def _split_stmt_rat(self, text: str):
        if not text: return "", ""
        parts = re.split(r"\bRationale:\s*", text, maxsplit=1, flags=re.I)
        stmt = parts[0].strip()
        rat  = parts[1].strip() if len(parts) > 1 else ""
        return stmt, rat

    def _requirement_text_fields(self, req_node: dict):
        raw = req_node.get("statement") or req_node.get("text") or req_node.get("description") or ""
        if isinstance(raw, dict) and raw.get("kind") == "markup":
            raw = raw.get("value","")
        raw = self._unwrap_markup(raw)
        if "statement" in req_node and "rationale" in req_node:
            return str(req_node.get("statement") or ""), str(req_node.get("rationale") or "")
        return self._split_stmt_rat(raw)

    def _requirement_relations(self) -> Dict[str, Set[str]]:
        of_interest = set(self._caplikes.keys()) | set(self._fcs.keys()) | set(self._funcs.keys()) | set(self._components.keys())
        req_to_objs: Dict[str, Set[str]] = {rid: set() for rid in self._requirements}
        # 1) direct ref scan on requirement nodes
        for rid, r in self._requirements.items():
            node = r["node"]
            for k in self.DEFAULT_REQ_REL_KEYS:
                if isinstance(node, dict) and k in node:
                    entries = node[k] if isinstance(node[k], list) else [node[k]]
                    for entry in entries:
                        oid = self._extract_id(entry if isinstance(entry, dict) else {"ref_uuid": entry})
                        if oid and oid in of_interest:
                            req_to_objs[rid].add(oid)
            for oid in self._iter_all_refs(node):
                if oid in of_interest: req_to_objs[rid].add(oid)

        # 2) via relation objects referenced in requirement['relations']
        for rid, r in self._requirements.items():
            node = r["node"]
            rels = node.get("relations")
            if not isinstance(rels, list): continue
            for entry in rels:
                link_id = None
                if isinstance(entry, dict):
                    link_id = entry.get("ref_uuid") or entry.get("uuid") or entry.get("id")
                if not link_id: continue
                rel_obj = self._relations_by_uuid.get(link_id)
                if not rel_obj: continue
                srcs = rel_obj.get("source") or []
                tgts = rel_obj.get("target") or []
                src_ids = [self._extract_id(s) for s in (srcs if isinstance(srcs, list) else [srcs])]
                tgt_ids = [self._extract_id(t) for t in (tgts if isinstance(tgts, list) else [tgts])]
                if rid in src_ids:
                    for oid in tgt_ids:
                        if oid in of_interest: req_to_objs[rid].add(oid)
                if rid in tgt_ids:
                    for oid in src_ids:
                        if oid in of_interest: req_to_objs[rid].add(oid)
        return req_to_objs


# CLI smoke
if __name__ == "__main__":
    yaml_file = "capella_model.yaml"
    rpt = CapellaYamlReport([yaml_file])
    print("Counts:", {"caps": sum(1 for v in rpt._caplikes.values() if v["kind"]=="Capability"),
                     "crs": sum(1 for v in rpt._caplikes.values() if v["kind"]=="CapabilityRealization"),
                     "fcs": len(rpt._fcs), "funcs": len(rpt._funcs),
                     "comps": len(rpt._components), "reqs": len(rpt._requirements)})

# ========================= DOCX generation & VCRM (v4) =========================

def _ensure_docx():
    """Ensure python-docx is available; if not, raise with a helpful message."""
    try:
        import docx  # noqa: F401
    except Exception as e:
        raise RuntimeError(
            "python-docx is required. Install with:\n  pip install python-docx"
        ) from e


def vcrm_dataframe(self) -> pd.DataFrame:
    """Build a Verification Cross-Reference Matrix (VCRM).

    Heuristic columns:
      - RequirementID, Requirement, Statement, Rationale
      - SuggestedMethod: based on where the requirement traces were found
        (Component→Inspection/Test, Function→Test/Analysis, FunctionalChain→Demonstration/Analysis, Direct→Inspection/Review)
      - Traceability: semicolon-separated 'Via: Object' entries
      - Capabilities: semicolon-separated capabilities associated to the requirement

    Notes
    -----
    This is a pragmatic VCRM for coursework; in production you would join to
    verification artifacts (test cases, analyses) via OSLC or tool APIs.
    """
    req_rel = self.requirements_relation_table()
    cap_req = self.capability_requirements_report()

    # Build per-requirement trace buckets from cap_req
    traces_by_req: Dict[str, List[Tuple[str, str]]] = {}
    caps_by_req: Dict[str, Set[str]] = {}

    for _, r in cap_req.iterrows():
        rid = str(r["RequirementID"])
        via = str(r["MatchedVia"])
        obj = str(r["MatchedObject"])
        traces_by_req.setdefault(rid, []).append((via, obj))
        caps_by_req.setdefault(rid, set()).add(str(r["Capability"]))

    # Method suggestion heuristic
    def suggest(vias: Set[str]) -> str:
        if "Component" in vias:
            return "Inspection / Test"
        if "Function" in vias:
            return "Test / Analysis"
        if "FunctionalChain" in vias:
            return "Demonstration / Analysis"
        if "Direct" in vias:
            return "Inspection / Review"
        return "TBD"

    rows = []
    seen = set()
    for _, rr in req_rel.iterrows():
        rid = str(rr["RequirementID"])
        if rid in seen:
            continue  # one row per requirement
        seen.add(rid)
        rname = str(rr["Requirement"])
        stmt  = str(rr.get("Statement", "") or "")
        rat   = str(rr.get("Rationale", "") or "")
        traces = traces_by_req.get(rid, [])
        vias = {v for v, _ in traces}
        method = suggest(vias)
        trace_str = "; ".join(sorted({f"{v}: {o}" for v, o in traces})) if traces else ""
        cap_str = "; ".join(sorted(caps_by_req.get(rid, set())))
        rows.append({
            "RequirementID": rid,
            "Requirement": rname,
            "Statement": stmt,
            "Rationale": rat,
            "SuggestedMethod": method,
            "Traceability": trace_str,
            "Capabilities": cap_str
        })

    df = pd.DataFrame.from_records(rows, columns=[
        "RequirementID","Requirement","Statement","Rationale","SuggestedMethod","Traceability","Capabilities"
    ])
    if not df.empty:
        df = df.sort_values(["Requirement"]).reset_index(drop=True)
    return df

CapellaYamlReport.vcrm_dataframe = vcrm_dataframe


def generate_requirements_docx(
    self,
    output_path: str,
    *,
    title: str = "System Requirements Specification",
    doc_number: str = "SRS-001",
    version: str = "0.1 (Draft)",
    organization: str = "SE Course",
) -> str:
    """Generate an IEEE-style SRS Word document with VCRM.

    The document contains:
      - Title page and Table of Contents
      - Introductory sections
      - Specific requirements grouped by Capability (with Statement/Rationale and trace table)
      - A VCRM section (Verification Cross-Reference Matrix)

    Parameters
    ----------
    output_path : str
        Destination .docx file path.
    title, doc_number, version, organization : str
        Title-page metadata.

    Returns
    -------
    str
        Absolute path to the generated DOCX.
    """
    _ensure_docx()
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import datetime

    # Build inputs
    cap_req = self.capability_requirements_report()
    req_rel = self.requirements_relation_table()
    vcrm_df = self.vcrm_dataframe()

    # Requirement metadata
    req_meta: Dict[str, Dict[str, str]] = {}
    for _, row in req_rel.iterrows():
        rid = str(row["RequirementID"])
        if rid not in req_meta:
            req_meta[rid] = {
                "Requirement": str(row["Requirement"]),
                "Statement": str(row.get("Statement", "") or ""),
                "Rationale": str(row.get("Rationale", "") or ""),
            }

    # Group by capability
    cap_groups: Dict[str, List[Dict[str, Any]]] = {}
    for _, row in cap_req.iterrows():
        cap = str(row["Capability"])
        cap_groups.setdefault(cap, []).append({k: row[k] for k in row.index})

    # Orphans
    all_req_ids = set(map(str, req_rel["RequirementID"].unique())) if len(req_rel) else set()
    req_ids_in_caps = set(map(str, cap_req["RequirementID"].unique())) if len(cap_req) else set()
    orphan_req_ids = sorted(all_req_ids - req_ids_in_caps)

    # ------------------ Build the DOCX ------------------
    doc = Document()

    # Title page
    title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title); run.bold = True; run.font.size = Pt(20)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{organization}\n").font.size = Pt(12)
    today = datetime.date.today().strftime("%B %d, %Y")
    sub.add_run(f"Document No.: {doc_number}\nVersion: {version}\nDate: {today}").font.size = Pt(12)
    doc.add_page_break()

    # Revision History (simple)
    doc.add_heading("Revision History", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Version"; hdr[1].text = "Date"; hdr[2].text = "Author"; hdr[3].text = "Summary of Changes"
    row = table.add_row().cells
    row[0].text = version; row[1].text = today; row[2].text = organization; row[3].text = "Initial draft generated from Capella YAML"
    doc.add_paragraph("")

    # TOC
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u'); p._p.append(fld)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1 Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph("This document specifies system requirements derived from the model export and generated automatically for consistency and traceability.")
    doc.add_heading("1.2 Scope", level=2); doc.add_paragraph("This specification covers functional and non-functional requirements for the system under study.")
    doc.add_heading("1.3 References", level=2); doc.add_paragraph("- IEEE 29148:2018 Requirements Engineering.")
    doc.add_heading("1.4 Definitions, Acronyms, and Abbreviations", level=2); doc.add_paragraph("TBD")
    doc.add_heading("1.5 Overview", level=2); doc.add_paragraph("Section 3 enumerates the specific requirements; Section 4 provides the Verification Cross-Reference Matrix (VCRM).")

    # 2. Overall Description (placeholders)
    doc.add_heading("2 Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2); doc.add_paragraph("TBD")
    doc.add_heading("2.2 Product Functions", level=2); doc.add_paragraph("TBD")
    doc.add_heading("2.3 User Classes and Characteristics", level=2); doc.add_paragraph("TBD")
    doc.add_heading("2.4 Operating Environment", level=2); doc.add_paragraph("TBD")
    doc.add_heading("2.5 Design and Implementation Constraints", level=2); doc.add_paragraph("TBD")
    doc.add_heading("2.6 Assumptions and Dependencies", level=2); doc.add_paragraph("TBD")

    # 3. Specific Requirements
    doc.add_heading("3 Specific Requirements", level=1)
    doc.add_paragraph("Requirements are grouped by Capability. Each includes the statement, rationale, and traceability to model elements.")

    def add_requirement_block(doc, rid: str, traces_df: pd.DataFrame):
        meta = req_meta.get(rid, {})
        name = meta.get("Requirement", "(Unnamed Requirement)")
        stmt = meta.get("Statement", "TBD")
        rat  = meta.get("Rationale", "")

        doc.add_heading(f"{name}  [{rid}]", level=3)
        p = doc.add_paragraph(); p.add_run("Statement: ").bold = True; p.add_run(str(stmt) or "TBD")
        if rat:
            p = doc.add_paragraph(); p.add_run("Rationale: ").bold = True; p.add_run(str(rat))

        if isinstance(traces_df, pd.DataFrame) and not traces_df.empty:
            # filter for this requirement if column present
            tdf = traces_df
            if "RequirementID" in tdf.columns:
                tdf = tdf[tdf["RequirementID"].astype(str) == rid]
            if not tdf.empty:
                t = doc.add_table(rows=1, cols=3); t.style = "Light List Accent 1"
                t.rows[0].cells[0].text = "Matched Via"
                t.rows[0].cells[1].text = "Model Element"
                t.rows[0].cells[2].text = "Element ID"
                seen = set()
                for _, r in tdf.iterrows():
                    key = (str(r.get("MatchedVia","")), str(r.get("MatchedObject","")), str(r.get("MatchedID","")))
                    if key in seen: continue
                    seen.add(key)
                    row = t.add_row().cells
                    row[0].text = key[0]; row[1].text = key[1]; row[2].text = key[2]

        doc.add_paragraph("Verification Method: TBD")
        doc.add_paragraph("Priority: TBD")
        doc.add_paragraph("Status: Draft")

    if cap_groups:
        for i, cap_name in enumerate(sorted(cap_groups.keys()), start=1):
            doc.add_heading(f"3.{i} Capability: {cap_name}", level=2)
            rows = cap_groups[cap_name]
            traces_df = pd.DataFrame(rows)
            for rid in sorted({str(r["RequirementID"]) for r in rows}):
                add_requirement_block(doc, rid, traces_df)

    # Orphans
    if orphan_req_ids:
        doc.add_heading("3.X Requirements Not Associated to a Capability", level=2)
        for rid in orphan_req_ids:
            add_requirement_block(doc, rid, pd.DataFrame(columns=["MatchedVia","MatchedObject","MatchedID"]))

    # 4. VCRM
    doc.add_heading("4 Verification Cross-Reference Matrix (VCRM)", level=1)
    doc.add_paragraph("This matrix lists each requirement with a suggested verification method and traceability context. Methods are heuristically suggested from model traces and should be reviewed.")
    v_cols = ["RequirementID","Requirement","SuggestedMethod","Capabilities","Traceability"]
    vdf = vcrm_df[v_cols] if not vcrm_df.empty else pd.DataFrame(columns=v_cols)
    t = doc.add_table(rows=1, cols=len(v_cols)); t.style = "Light List Accent 1"
    for i, c in enumerate(v_cols):
        t.rows[0].cells[i].text = c
    for _, r in vdf.iterrows():
        row = t.add_row().cells
        for i, c in enumerate(v_cols):
            row[i].text = str(r.get(c, ""))

    # Footer note
    doc.add_paragraph("Note: Use Word’s References → Update Table to refresh the Table of Contents.")

    abspath = os.path.abspath(output_path)
    doc.save(abspath)
    return abspath

CapellaYamlReport.generate_requirements_docx = generate_requirements_docx
